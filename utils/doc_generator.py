from io import BytesIO
from docx import Document

def create_cover_letter_docx(cover_letter_text,candidate_name="Candidate"):
    doc = Document()

    doc.add_heading(f"{candidate_name}-Cover Letter",level=1)
    doc.add_paragraph(cover_letter_text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


