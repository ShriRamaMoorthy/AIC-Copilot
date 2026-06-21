import os
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings

BASE_DIR = os.path.dirname(
    os.path.dirname(
        os.path.abspath(__file__)
    )
)

KB_PATH = os.path.join(BASE_DIR,"knowledge_base")
CHROMA_PATH = os.path.join(BASE_DIR,"vector_db","chroma_db")

def load_documents():
    documents = []

    for root, _, files in os.path(KB_PATH):
        folder_name = os.path.basename(root)
        for file in files:
            file_path = os.path.join(root,file)
            text = ""

            try:
                if file.endswith(".txt"):
                    with open(file_path,'r',encoding="utf-8") as f:
                        text = f.read()
                elif file.endswith(".pdf"):
                    reader = PdfReader(file_path)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text+=page_text+"\n"
                elif file.endswith(".docx"):
                    doc = DocxDocument(file_path)
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text+=(para.text+'\n')
                
                if text.strip():
                    documents.append({
                        "content":text,
                        "source":file,
                        "category":folder_name
                    })
            
            except Exception as e:
                print(f'Error reading {file_path}:{e}')
    return documents


def chunk_documents(documents):
    splitter = (RecursiveCharacterTextSplitter(
        chunk_size=500,chunk_overlap=100
    ))
    chunks=[]
    for doc in documents:
        split_texts = splitter.split_text(
            doc["content"]
        )

        for chunk in split_texts:
            chunks.append(

                Document(
                    page_content=chunk,

                metadata={
                    "source":doc["source"],
                    "category":doc["category"]
                }
                )
            )
    return chunks


def build_vector_db():
    docs = load_documents()
    chunks = chunk_documents(docs)
    print(f'Loaded {len(chunks)} chunks')

    if not chunks:
        raise ValueError("No Chunks Found.")
    
    embeddings = (HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2"))

    if os.path.exists(CHROMA_PATH):
        import shutil

        shutil.rmtree(CHROMA_PATH)
    
    vector_db = (
        Chroma.from_documents(
            documents=chunks,
            embedding=embeddings,
            persist_directory=CHROMA_PATH
        )
    )

    print("Vector DB created successfully")

    if __name__ == "__main__":
        build_vector_db()